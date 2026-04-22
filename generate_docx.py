from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(9.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(1)
style.paragraph_format.line_spacing = 1.1

blue = RGBColor(0x2a, 0x5a, 0xa6)
black = RGBColor(0x1a, 0x1a, 0x1a)
gray = RGBColor(0x55, 0x55, 0x55)
dark = RGBColor(0x33, 0x33, 0x33)

# ---- Helper functions ----
def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {qn('w:val'): 'single', qn('w:sz'): '4', qn('w:space'): '1', qn('w:color'): '1a1a1a'})
    pBdr.append(bottom)
    pPr.append(pBdr)

def section_title(text):
    add_divider()
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = black
    run.font.name = 'Cambria'
    p.paragraph_format.space_after = Pt(4)

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.1
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = black
        r2 = p.add_run(text)
        r2.font.size = Pt(9)
        r2.font.color.rgb = dark
    else:
        r = p.add_run(text)
        r.font.size = Pt(9)
        r.font.color.rgb = dark

def exp_header(title, date):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = black
    r2 = p.add_run(f'    {date}')
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = gray
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(3)

def exp_org(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = gray
    p.paragraph_format.space_after = Pt(2)

def add_text(text, size=9.5, bold=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color or dark
    p.paragraph_format.space_after = Pt(1)
    return p

# ============ HEADER ============
p = doc.add_paragraph()
r = p.add_run('Kayoz Mistry')
r.font.size = Pt(24)
r.font.bold = True
r.font.color.rgb = black
r.font.name = 'Cambria'
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run('Head Media Technology | BMS & RMS Platforms | ERP & AdTech Systems | AI-Powered Automation | Cloud Migration (AWS/Azure) | Digital Transformation | CSPO\u00AE')
r.font.size = Pt(8.5)
r.font.color.rgb = blue
r.font.bold = True
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run('Mumbai, India  |  +91-9819673133  |  kayozm@gmail.com  |  linkedin.com/in/kayozmistry')
r.font.size = Pt(8.5)
r.font.color.rgb = gray
p.paragraph_format.space_after = Pt(4)

# ============ CORE COMPETENCIES ============
p = doc.add_paragraph()
kw = 'BMS \u2022 RMS \u2022 ERP Implementation \u2022 AdTech & Ad Sales Scheduling \u2022 MAM Integration \u2022 Linear Playout \u2022 OTT & FAST Channels \u2022 Cloud Migration (AWS/Azure) \u2022 AI & Automation \u2022 Digital Transformation \u2022 SSAI/CSAI \u2022 SCTE-35 \u2022 HA/DR Architecture \u2022 SAP & ERP Integration \u2022 Product Strategy \u2022 P&L Awareness \u2022 Vendor Management \u2022 Team Leadership \u2022 Stakeholder Mgmt (C-Suite) \u2022 Pre-Sales \u2022 CSPO\u00AE'
r = p.add_run(kw)
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
p.paragraph_format.space_after = Pt(3)

# ============ EXECUTIVE SUMMARY ============
section_title('Executive Summary')
p = doc.add_paragraph()
parts = [
    ('Media technology leader with ', False), ('14+ years', True),
    (' driving enterprise-grade broadcast implementations, ERP system rollouts, AdTech platform deployments, and AI-powered automation across ', False),
    ('four continents', True), ('. Deep domain expertise in ', False),
    ('Broadcast Management Systems (BMS)', True), (', ', False),
    ('Rights Management Systems (RMS)', True), (', ', False),
    ('ERP integration', True),
    (', Ad Sales traffic & scheduling platforms, and MAM integrations. Led global deployments for ', False),
    ('Sony Pictures Networks, Viacom18, SABC, Warner Bros, Tata Play, Network18, and Econet Media', True),
    (' across India, UAE, South Africa, and Asia-Pacific. Built and scaled a media technology function from the ground up over 14 years. Proven track record implementing scalable solutions that ', False),
    ('reduced operational costs by 40%', True), (' and ', False),
    ('increased efficiency by 30%', True),
    (' through AI and automation. CSPO-certified with deep expertise in product strategy, stakeholder management, and end-to-end system integration.', False),
]
for text, bold in parts:
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.bold = bold
    r.font.color.rgb = black if bold else dark
p.paragraph_format.space_after = Pt(4)

for bp, rest in [
    ('40% operational cost reduction', ' through automation, cloud migration, and process re-engineering'),
    ('30% efficiency gain', ' via AI-powered workflows, chatbots, self-service portals, and automated scheduling'),
    ('99.9% broadcast uptime', ' achieved through HA/DR architecture and proactive monitoring'),
    ('10+ successful broadcast deployments', ' across India, UAE, South Africa, Papua New Guinea, and Asia-Pacific'),
    ('Fortune 500 & Tier-1 clients', ' including Sony, Warner Bros, Viacom18, SABC, Tata Play, Network18, and Reliance'),
    ('14-year growth trajectory', ' \u2014 Executive \u2192 Manager \u2192 Senior Manager \u2192 Head of Media Technology'),
]:
    add_bullet(rest, bold_prefix=bp)

# ============ EXPERIENCE ============
section_title('Professional Experience')

exp_header('Head Media Technology', 'April 2025 \u2013 Present')
exp_org('UTO Solutions, Mumbai, India')
for b in [
    'Lead and scale a cross-functional team of engineers, implementation specialists, and support staff \u2014 setting technical direction, defining performance KPIs, and building a culture of ownership and accountability',
    'Spearhead media technology strategy with focus on BMS, RMS, and ERP platforms \u2014 aligning Ad Sales scheduling, traffic management, rights clearance, and ERP-driven revenue workflows',
    'Own the technology vision, roadmap, and budget \u2014 presenting quarterly strategy updates to the Board of Directors and aligning CAPEX/OPEX investment with client and market needs',
    'Architect next-generation broadcast-to-digital workflows integrating BMS and ERP with OTT scheduling, FAST channels, programmatic ad insertion (SSAI/CSAI), and dynamic ad decisioning',
    'Drive AI and automation initiatives modernizing legacy BMS/RMS/ERP systems \u2014 reducing manual scheduling and accelerating campaign turnaround',
    'Oversee ERP integration strategy connecting SAP and media-specific ERP modules with BMS for finance, invoicing, contract management, amortization, and revenue assurance',
    'Champion talent development and succession planning \u2014 structured onboarding, upskilling programs, and cross-training rotations',
    'Serve as executive point of contact for strategic client relationships \u2014 C-level business reviews, contract negotiations, and partnership growth',
    'Define product roadmap aligning BMS/ERP capabilities with AdTech demands: CTV, addressable advertising, hybrid monetization (AVOD/SVOD/FAST), first-party data strategies',
]:
    add_bullet(b)

exp_header('Senior Manager \u2013 Client Services & Implementation', '2025 \u2013 April 2025')
exp_org('UTO Solutions, Mumbai, India')
for b in [
    'Led pre-sales for BMS and RMS platforms \u2014 product demos, solution architecture workshops, and RFP responses for Tier-1 broadcasters',
    'Implemented AI-powered workflows automating support queries (+30% response speed); developed AI chatbot reducing manual escalations',
    'Launched self-service client portal with AI-driven search enabling independent issue resolution',
    'Managed end-to-end BMS/RMS and ERP implementations across India, UAE, South Africa, and Papua New Guinea including SAP-to-BMS integration',
    'Created monthly performance reports for Board of Directors enabling strategic decision-making',
    'Served as single point of contact for key accounts managing critical issues, SLA compliance, and long-term relationships',
    'Led change management during BMS/RMS system transitions ensuring minimal disruption to on-air schedules',
]:
    add_bullet(b)

exp_header('Manager \u2013 Client Service & Implementation', '2017 \u2013 2025')
exp_org('UTO Solutions, Mumbai, India')
for b in [
    'Led automation strategies reducing manual workload by 25% and processing time by 40% across BMS scheduling and traffic operations',
    'Implemented HA and DR solutions for mission-critical BMS/RMS software ensuring 99.9% uptime and broadcast continuity',
    'Managed system integrations using APIs to bridge BMS with ERP (SAP), MAM, playout, finance, invoicing, and third-party AdTech platforms',
    'Subject matter expert advising Sony Pictures Networks, Viacom18, Tata Play, Warner Bros, SABC, and Network18 on broadcast scheduling, traffic, and rights clearance',
    'Created comprehensive BRDs, SOPs, and training documentation for clients and internal teams',
    'Led pre-sales: product demonstrations, technical proposals, and RFP responses for BMS/RMS platform expansion',
    'Managed IT infrastructure including firewall, network routing, and server architecture for hosted BMS deployments',
    'Collaborated with Product, QA, and Engineering teams incorporating client feedback into product roadmap',
]:
    add_bullet(b)

exp_header('Assistant Manager \u2013 Client Services & Implementation', '2016 \u2013 2017')
exp_org('UTO Solutions, Mumbai, India')
for b in [
    'Delivered onsite BMS/RMS training and software implementations with comprehensive training materials and knowledge transfer',
    'Conducted UAT and validated post-implementation functionality for broadcast scheduling and ad traffic modules',
    'Led change management gathering requirements during system updates and BMS version migrations',
    'Created wireframes and prototypes using Indigo Studio and Axure for proposed BMS UI enhancements',
]:
    add_bullet(b)

exp_header('Client Service Executive', '2011 \u2013 2016')
exp_org('UTO Solutions, Mumbai, India')
for b in [
    'Managed full SDLC for custom BMS solutions including requirements analysis, functional specs, and documentation',
    'Conducted application testing and QA for broadcast scheduling, traffic management, and rights management modules',
    'Tracked and resolved issues using JIRA and ZOHO ensuring timely resolution across multiple broadcast clients',
    'Delivered onsite training and implementation for BMS platforms across diverse broadcast environments',
]:
    add_bullet(b)

# ============ GLOBAL IMPLEMENTATIONS ============
section_title('Global Implementations (10+ Successful Broadcast Deployments)')

implementations = [
    ('SABC', '2024', 'Johannesburg, South Africa'),
    ('Sentech Media / FreeVision', '2024', 'South Africa'),
    ('Econet Media / Kwese TV', '2016\u201317', 'Johannesburg, South Africa'),
    ('TataSky / Tata Play', '2015', 'Delhi & Mumbai, India'),
    ('Moby Media Group', '2015', 'Dubai, UAE'),
    ('Alliance Media / Urdu1', '2015', 'Dubai, UAE'),
    ('Image Nation / Quest Arabia', '2015', 'Abu Dhabi, UAE'),
    ('Network18', '2013', 'Mumbai, India'),
    ('Reliance Big Magic / Big CBS', '2013', 'Surat & Mumbai, India'),
    ('EMTV', '2013', 'Port Moresby, Papua New Guinea'),
    ('Sony Entertainment Television', '2012', 'Mumbai, India'),
    ('Viacom18', '2012', 'Mumbai, India'),
]
for name, year, loc in implementations:
    add_bullet(f' ({year}) \u2014 {loc}', bold_prefix=name)

# ============ KEY PROJECTS & SOLUTIONS ============
section_title('Key Projects & Solutions Delivered')

for bp, rest in [
    ('SABC BroadView Implementation', ' \u2013 Senior SME & Project Manager migrating channels to BroadView BMS (Dec 2024)'),
    ('FreeVision OTT Platform Launch', ' \u2013 PM for Sentech Media OTT across web, TV, and mobile with BMS-driven scheduling (Apr 2024)'),
    ('BroadView AWS Cloud Migration', ' \u2013 Project Lead migrating BMS to AWS improving uptime and DR (Feb 2023)'),
    ('Sony Pictures Networks BroadView', ' \u2013 Senior SME & PM streamlining ad scheduling and metadata management (Jun 2012)'),
]:
    add_bullet(rest, bold_prefix=bp)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
r = p.add_run('Technical Solutions Built & Delivered:')
r.font.size = Pt(9)
r.font.bold = True
r.font.color.rgb = black
p.paragraph_format.space_after = Pt(2)

solutions = [
    'Broadcast Scheduling & Traffic (Ad Sales) Software',
    'TRAI Violation Reporting (Web Service)',
    'EPG Generator for DTH Platforms (Web Application)',
    'SAP Integration Web Application',
    'Amortization & Revenue Reporting Module',
    'Invoicing Application (PDF Invoice Generation)',
    'Geo-Targeting & Geo-Advertising Implementations',
    'On Demand Scheduler Module (OTT)',
    'MAM Automation Integration (AVID, DALET, Flex)',
    'SCTE-35 Local Ad Insertion on Linear Platform',
    'Secondary Element Overlay Automation (OnScreen Graphics)',
    'Music & Radio Scheduling Software',
    'Secondary Event ID Tracker',
    'Target vs Achievement Sales Report Analytics',
    'BroadView-Pebble Playout Integration',
    'Content & Rights Migration Tools',
    'BondFlex MAM Integration',
    'RMS Rights Clearance Module',
]
for s in solutions:
    add_bullet(s)

# ============ TECHNICAL SKILLS ============
section_title('Technical Skills')
skills = [
    ('AdTech & BMS', 'BMS, RMS, Ad Sales Traffic & Scheduling, BroadView, EPG Generation, SSAI/CSAI, SCTE-35 Local Ad Insertion, FAST Channel Scheduling, Geo-Targeting & Geo-Advertising, CTV Advertising, TRAI Compliance & Reporting, Amortization & Revenue Assurance, AVOD/SVOD/FAST, Music & Radio Scheduling'),
    ('ERP & Enterprise', 'ERP Implementation & Integration, SAP (Finance, Invoicing, Revenue), Media-Specific ERP Modules, Contract & Vendor Management, Finance & Billing Integration, Amortization Reporting, BI & Analytics Dashboards'),
    ('Broadcast', 'MAM (Dalet, Flex, Avid, Xytech), Playout (Imagine, Pebble Beach, Harmonic, iTX, GrassValley), Linear TV & Radio, OTT Platforms & Player, OTT Streaming/CDN/SSAI, DRM, Master Control Automation, OnScreen Graphics/Overlay Automation, Adobe After Effects Integration'),
    ('Cloud & Infra', 'AWS (Elemental, MediaLive), Microsoft Azure, HA/DR Planning & Execution, Cloud Migration (CAPEX-to-OPEX), Server Architecture, Network Routing & Firewall, IT Infrastructure & Hardware'),
    ('Programming', 'SQL, Firebird, Shell Scripting, PowerShell, REST API Integration, XML, JSON, Web Services, Database Migration'),
    ('PM & Design', 'JIRA, Asana, ClickUp, Agile/Scrum (CSPO\u00AE), Figma, Balsamiq, Adobe XD, Lucidchart, Visio, Draw.io, Indigo Studio, Axure'),
    ('Leadership', 'Product Strategy & Roadmap, P&L Awareness, Pre-Sales & Solution Architecture, Business Analysis (BRD/SOP/SOW), Vendor Sourcing & Contract Negotiation, Cross-Functional Team Building, Change Management, Customer Success, Stakeholder Management (C-Suite), Digital Transformation'),
]
for cat, val in skills:
    p = doc.add_paragraph()
    r = p.add_run(f'{cat}: ')
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = black
    r2 = p.add_run(val)
    r2.font.size = Pt(9)
    r2.font.color.rgb = dark
    p.paragraph_format.space_after = Pt(2)

# ============ EDUCATION ============
section_title('Education & Certification')
for name, org, date in [
    ('Certified Scrum Product Owner (CSPO\u00AE)', 'Scrum Alliance', 'February 2023'),
    ('Diploma in Software Engineering', 'NIIT Mumbai (GNIIT Certified)', 'April 2011'),
    ('Bachelor of Commerce', 'Mumbai University, Thakur College of Engineering & Technology', 'April 2011'),
    ('Higher Secondary Education', 'Lords Junior College, Mumbai University', '2008'),
]:
    p = doc.add_paragraph()
    r = p.add_run(name)
    r.font.size = Pt(9.5)
    r.font.bold = True
    r.font.color.rgb = black
    r2 = p.add_run(f'  \u2014  {org}')
    r2.font.size = Pt(9)
    r2.font.italic = True
    r2.font.color.rgb = gray
    r3 = p.add_run(f'    {date}')
    r3.font.size = Pt(8.5)
    r3.font.color.rgb = gray
    p.paragraph_format.space_after = Pt(2)

# ============ LANGUAGES ============
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
r = p.add_run('Languages: ')
r.font.size = Pt(9)
r.font.bold = True
r.font.color.rgb = black
r2 = p.add_run('English (Fluent) \u2022 Hindi (Native) \u2022 Gujarati (Native) \u2022 Marathi (Conversational)')
r2.font.size = Pt(9)
r2.font.color.rgb = dark

p = doc.add_paragraph()
r = p.add_run('Industry Focus: ')
r.font.size = Pt(9)
r.font.bold = True
r.font.color.rgb = black
r2 = p.add_run('Media & Broadcasting \u2022 AdTech \u2022 OTT/Streaming \u2022 Enterprise Software')
r2.font.size = Pt(9)
r2.font.color.rgb = dark

doc.save(r'C:\Users\utoadmin\portfolio-variants\Kayoz_Mistry_Resume.docx')
print('DOCX saved successfully')
